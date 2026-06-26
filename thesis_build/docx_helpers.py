# -*- coding: utf-8 -*-
"""
Yalova Üniversitesi bitirme tezi şablonu üzerinde python-docx ile içerik üretmek
için yardımcı fonksiyonlar. Stiller şablonun stil ID'leriyle (Balk1..6, Normal,
ekilYazs, izelgeMetni, ListeMaddemi, DaraltlmMetin, Dipnot) uygulanır.
"""
import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE = os.path.dirname(__file__)
FIG = os.path.join(HERE, "figures")


# ── stil uygulama ─────────────────────────────────────────────────────────────
def set_style(p, style_id):
    """Paragrafa şablon stil ID'siyle stil uygular (ada göre değil)."""
    pPr = p._p.get_or_add_pPr()
    for e in pPr.findall(qn("w:pStyle")):
        pPr.remove(e)
    ps = OxmlElement("w:pStyle")
    ps.set(qn("w:val"), style_id)
    pPr.insert(0, ps)
    return p


# ── alan (field) ekleme ───────────────────────────────────────────────────────
def add_field(paragraph, instr, placeholder="»alanı güncelleyin (F9)«", dirty=True):
    r = paragraph.add_run()
    b = OxmlElement("w:fldChar")
    b.set(qn("w:fldCharType"), "begin")
    if dirty:
        b.set(qn("w:dirty"), "true")
    r._r.append(b)

    r2 = paragraph.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = " " + instr + " "
    r2._r.append(it)

    r3 = paragraph.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r3._r.append(sep)

    paragraph.add_run(placeholder)

    r5 = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r5._r.append(end)


# ── sayfa numarası biçimi ──────────────────────────────────────────────────────
def set_page_number_format(section, fmt, start=None):
    """sectPr içine pgNumType ekler (fmt: 'lowerRoman' | 'decimal')."""
    sectPr = section._sectPr
    pg = sectPr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sectPr.append(pg)
    pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))


def clear_hf_refs(section):
    """sectPr'dan tüm header/footer referanslarını (şablondan miras kalanlar) siler."""
    sectPr = section._sectPr
    for tag in ("w:headerReference", "w:footerReference"):
        for e in sectPr.findall(qn(tag)):
            sectPr.remove(e)


def clear_pgnum(section):
    sectPr = section._sectPr
    pg = sectPr.find(qn("w:pgNumType"))
    if pg is not None:
        sectPr.remove(pg)


def put_page_field(container, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Bir header/footer'a ortalı PAGE alanı koyar."""
    p = container.paragraphs[0]
    p.alignment = alignment
    add_field(p, "PAGE", placeholder="1", dirty=False)


# ── tablo kenarlıkları ─────────────────────────────────────────────────────────
def set_table_borders(table, sz=6):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "000000")
        borders.append(e)
    tblPr.append(borders)


def shade_cell(cell, fill="D9D9D9"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


# ── içerik yazıcı sınıfı ───────────────────────────────────────────────────────
class Thesis:
    def __init__(self, base_path):
        self.doc = Document(base_path)
        # base.docx gövdesi boş (yalnız sectPr). İlk içerik buraya yazılır.

    # --- temel bloklar ---
    @staticmethod
    def _page_break_before(p):
        pPr = p._p.get_or_add_pPr()
        pPr.append(OxmlElement("w:pageBreakBefore"))

    def h(self, level, text, brk=False):
        p = self.doc.add_paragraph()
        set_style(p, f"Balk{level}")
        if brk:
            self._page_break_before(p)
        p.add_run(text)
        return p

    def h5(self, text, brk=True):
        p = self.doc.add_paragraph()
        set_style(p, "Balk5")
        if brk:
            self._page_break_before(p)
        p.add_run(text)
        return p

    def h6(self, text, brk=True):
        p = self.doc.add_paragraph()
        set_style(p, "Balk6")
        if brk:
            self._page_break_before(p)
        p.add_run(text)
        return p

    def para(self, text, style="Normal"):
        p = self.doc.add_paragraph()
        set_style(p, style)
        p.add_run(text)
        return p

    def bullet(self, text, level=1):
        p = self.doc.add_paragraph()
        set_style(p, "ListeMaddemi" if level == 1 else f"ListeMaddemi{level}")
        p.add_run(text)
        return p

    def title_block(self, lines, sizes=None, bold=True, space_before=0):
        """Kapak için ortalı, koyu satır bloğu."""
        for i, line in enumerate(lines):
            p = self.doc.add_paragraph()
            set_style(p, "Normal")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line)
            r.bold = bold
            if sizes and i < len(sizes) and sizes[i]:
                r.font.size = Pt(sizes[i])
        return

    def spacer(self, n=1):
        for _ in range(n):
            p = self.doc.add_paragraph()
            set_style(p, "Normal")

    def page_break(self):
        from docx.enum.text import WD_BREAK
        p = self.doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)

    # --- şekil ---
    def figure(self, filename, caption, width_cm=15.0):
        path = os.path.join(FIG, filename)
        p = self.doc.add_paragraph()
        set_style(p, "Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Cm(width_cm))
        cap = self.doc.add_paragraph()
        set_style(cap, "ekilYazs")
        cap.add_run(caption)
        return

    # --- çizelge ---
    def table_caption(self, caption):
        cap = self.doc.add_paragraph()
        set_style(cap, "izelgeMetni")
        cap.add_run(caption)

    def table(self, rows, header=True, widths_cm=None, font_size=10):
        ncol = len(rows[0])
        t = self.doc.add_table(rows=len(rows), cols=ncol)
        t.alignment = 1  # center
        set_table_borders(t)
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = t.cell(ri, ci)
                cell.text = ""
                para = cell.paragraphs[0]
                set_style(para, "Normal")
                # kompakt hücre: tek aralık
                pPr = para._p.get_or_add_pPr()
                sp = OxmlElement("w:spacing")
                sp.set(qn("w:before"), "20")
                sp.set(qn("w:after"), "20")
                sp.set(qn("w:line"), "240")
                sp.set(qn("w:lineRule"), "auto")
                pPr.append(sp)
                run = para.add_run(str(val))
                run.font.size = Pt(font_size)
                if header and ri == 0:
                    run.bold = True
                    shade_cell(cell, "E7E7E7")
                if widths_cm:
                    cell.width = Cm(widths_cm[ci])
        return t

    def save(self, path):
        self.doc.save(path)
