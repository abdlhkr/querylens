# -*- coding: utf-8 -*-
"""
QueryLens bitirme tezini (.docx) Yalova Üniversitesi şablonuna göre üretir.
Çalıştır: python3 thesis_build/gen_thesis.py
Çıktı:    thesis_build/QueryLens_Bitirme_Tezi.docx
"""
import os
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from docx_helpers import (
    Thesis, set_page_number_format, put_page_field, clear_hf_refs, clear_pgnum,
)

HERE = os.path.dirname(__file__)
BASE = os.path.join(HERE, "base.docx")
OUT = os.path.join(HERE, "QueryLens_Bitirme_Tezi.docx")

T = Thesis(BASE)
P = T.para
B = T.bullet

# ── Kapak bilgileri (kullanıcı tarafından doldurulan değerler) ────────────────
BOLUM = "BİLGİSAYAR MÜHENDİSLİĞİ BÖLÜMÜ"
# Tez başlığı (kapakta birden çok satıra bölünebilir)
BASLIK = [
    "GERÇEK ZAMANLI YÜZ İFADESİ DUYGU ANALİZİ",
    "VE İÇERİK ÖNERİ SİSTEMİ",
]
OGRENCI = "Abdullah Kara"
OGRENCI_NO = "Öğrenci No"
DANISMAN = "Dr. Öğr. Üyesi Murat Okkalıoğlu"
DONEM = "2025 – 2026 Bahar Yarıyılı"


# ══════════════════════════════════════════════════════════════════════════════
# DIŞ KAPAK  (Ek-2 formatı)
# ══════════════════════════════════════════════════════════════════════════════
T.spacer(2)
T.title_block(["YALOVA ÜNİVERSİTESİ", "MÜHENDİSLİK FAKÜLTESİ"], sizes=[16, 16])
T.title_block([BOLUM], sizes=[14])
T.spacer(2)
T.title_block(["- BİTİRME TEZİ -"], sizes=[14])
T.spacer(2)
T.title_block(BASLIK, sizes=[16] * len(BASLIK))
T.spacer(3)
T.title_block([OGRENCI], sizes=[13], bold=False)
T.spacer(2)
T.title_block([f"Bitirme Tezi Danışmanı: {DANISMAN}"], sizes=[12], bold=False)
T.spacer(3)
T.title_block(["YALOVA, 2026"], sizes=[13])

# ── İÇ KAPAK (Ek-3 formatı) ───────────────────────────────────────────────────
T.page_break()
T.spacer(2)
T.title_block(["YALOVA ÜNİVERSİTESİ", "MÜHENDİSLİK FAKÜLTESİ"], sizes=[16, 16])
T.title_block([BOLUM], sizes=[14])
T.spacer(2)
T.title_block(BASLIK, sizes=[15] * len(BASLIK))
T.spacer(2)
T.title_block([OGRENCI, OGRENCI_NO], sizes=[12, 12], bold=False)
T.spacer(2)
T.title_block([
    f"1. Bitirme Tezi Danışmanı: {DANISMAN}",
    "2. Jüri Üyesi:",
    "3. Jüri Üyesi:",
], sizes=[12, 12, 12], bold=False)
T.spacer(2)
T.title_block([f"Bitirme Tezinin Dönemi: {DONEM}"], sizes=[12], bold=False)


# ══════════════════════════════════════════════════════════════════════════════
# ÖN SAYFALAR BÖLÜMÜ (roma rakamı)
# ══════════════════════════════════════════════════════════════════════════════
T.doc.add_section(WD_SECTION.NEW_PAGE)

# ── İÇİNDEKİLER ───────────────────────────────────────────────────────────────
h = T.h5("İÇİNDEKİLER", brk=False)
note = T.doc.add_paragraph()
from docx_helpers import set_style
set_style(note, "Normal")
add_p = note.add_run("Aşağıdaki listeler otomatik alanlardır. Microsoft Word'de açtıktan sonra "
                     "tüm belgeyi seçip F9 tuşuna basarak (veya alana sağ tıklayıp "
                     "“Alanı Güncelleştir”) içindekiler, şekil ve çizelge listeleri ile "
                     "sayfa numaralarını güncelleyiniz.")
add_p.italic = True
toc = T.doc.add_paragraph()
set_style(toc, "Normal")
from docx_helpers import add_field
add_field(toc, 'TOC \\o "1-6" \\h \\z \\u')

# ── SİMGE LİSTESİ ─────────────────────────────────────────────────────────────
T.h5("SİMGE LİSTESİ")
P("Bu tezde özel matematiksel simge kullanılmamıştır. Kullanılan tek nicel ifade, "
  "bir tablonun önem (importance) katsayısının hesaplanmasında geçen logaritmik "
  "ifadedir:")
T.para("importance = log₁₀(satır_sayısı + 1)", style="DaraltlmMetin")
P("Bu katsayı, getirim aşamasında aday tabloların yeniden sıralanmasında kullanılan "
  "boyutsuz bir büyüklüktür (bkz. Bölüm 4.3).")

# ── KISALTMA LİSTESİ ──────────────────────────────────────────────────────────
T.h5("KISALTMA LİSTESİ")
abbr = [
    ("API", "Application Programming Interface (Uygulama Programlama Arayüzü)"),
    ("BM25", "Best Matching 25 (kelime tabanlı benzerlik sıralama fonksiyonu)"),
    ("CI/CD", "Continuous Integration / Continuous Delivery"),
    ("DDL", "Data Definition Language (Veri Tanımlama Dili)"),
    ("DML", "Data Manipulation Language (Veri İşleme Dili)"),
    ("ER", "Entity–Relationship (Varlık–İlişki)"),
    ("gRPC", "Google Remote Procedure Call"),
    ("HTTP", "Hypertext Transfer Protocol"),
    ("JSON", "JavaScript Object Notation"),
    ("JWT", "JSON Web Token"),
    ("LLM", "Large Language Model (Büyük Dil Modeli)"),
    ("ORM", "Object–Relational Mapping"),
    ("OTP", "One-Time Password (Tek Kullanımlık Şifre)"),
    ("RAG", "Retrieval-Augmented Generation (Getirim-Destekli Üretim)"),
    ("RASL", "Retrieval-Augmented Schema Linking"),
    ("REST", "Representational State Transfer"),
    ("SaaS", "Software as a Service (Hizmet Olarak Yazılım)"),
    ("SDK", "Software Development Kit"),
    ("SQL", "Structured Query Language"),
    ("UUID", "Universally Unique Identifier"),
    ("VTYS", "Veri Tabanı Yönetim Sistemi"),
]
for k, v in abbr:
    p = T.para(f"{k}\t{v}", style="ListeMetni")

# ── ŞEKİL LİSTESİ ─────────────────────────────────────────────────────────────
T.h5("ŞEKİL LİSTESİ")
fl = T.doc.add_paragraph()
set_style(fl, "Normal")
add_field(fl, 'TOC \\h \\z \\t "Şekil Yazısı;1"')

# ── ÇİZELGE LİSTESİ ───────────────────────────────────────────────────────────
T.h5("ÇİZELGE LİSTESİ")
cl = T.doc.add_paragraph()
set_style(cl, "Normal")
add_field(cl, 'TOC \\h \\z \\t "Çizelge Metni;1"')

# ── ÖNSÖZ ─────────────────────────────────────────────────────────────────────
T.h5("ÖNSÖZ")
P("Veri, çağımızın en değerli kaynaklarından biri hâline gelmiştir; ancak bu verinin "
  "değeri, ancak ona erişebilen ve onu yorumlayabilen kişilerce ortaya çıkarılabilir. "
  "Bugün kurumların büyük bir kısmında veriler ilişkisel veri tabanlarında saklanmakta, "
  "bu verilere erişim ise çoğunlukla SQL bilgisi gerektirmektedir. Bu durum, teknik "
  "olmayan kullanıcılar ile veri arasında görünmez bir duvar oluşturmaktadır. Bu tez, "
  "söz konusu duvarı doğal dil ile veri tabanı arasında bir köprü kurarak aşmayı "
  "amaçlayan QueryLens platformunun tasarımını ve gerçekleştirimini konu almaktadır.")
P("Çalışmanın özünü, devasa veri tabanı şemalarında dahi doğru SQL üretebilmek için "
  "geliştirilen getirim-destekli şema bağlama yaklaşımı oluşturmaktadır. Bu yaklaşım, "
  "Amazon araştırmacılarının önerdiği RASL (Retrieval-Augmented Schema Linking) "
  "yönteminden ilham almış ve mikroservis tabanlı, gerçek kullanıma uygun bir SaaS "
  "platformu içerisinde uçtan uca hayata geçirilmiştir.")
P(f"Tez çalışmam boyunca değerli yönlendirmeleriyle bana destek olan danışmanım Sayın "
  f"{DANISMAN}’na, bilgi ve deneyimlerini benimle paylaşan bölüm hocalarıma ve "
  "her aşamada yanımda olan aileme teşekkürlerimi sunarım.")

# ── ÖZET ──────────────────────────────────────────────────────────────────────
T.h5("ÖZET")
P("Bu tez, kullanıcıların kendi ilişkisel veri tabanlarını bir aracı (agent) "
  "üzerinden bağlayıp doğal dil ile sorgulayabildiği QueryLens adlı bir SaaS "
  "platformunu konu almaktadır. Platform; React tabanlı bir ön yüz, Spring Cloud "
  "Gateway, kimlik doğrulama, kullanıcı, veri tabanı ve yapay zeka servislerinden "
  "oluşan bir mikroservis mimarisi ile kullanıcının veri tabanında çalışan bir "
  "WebSocket aracısından meydana gelmektedir. Çalışmanın merkezinde, büyük dil "
  "modelleriyle doğal dilden SQL üretiminin en önemli ölçeklenme sorunu yer "
  "almaktadır: gerçek kurumsal veri tabanları yüzlerce tablo içerebildiğinden, tüm "
  "şemanın modele verilmesi hem bağlam sınırlarını aşmakta hem de doğruluğu "
  "düşürmektedir. Bu sorun, getirim-destekli şema bağlama (RASL) yaklaşımının bir "
  "uyarlaması ile çözülmüştür. Veri tabanı doğrulandığında sistem otomatik bir içe "
  "bakış (introspection) yürüterek her tablonun kolon, tip, birincil/yabancı anahtar "
  "ve yaklaşık satır sayısı bilgisini çıkarır; her tablo için büyük dil modeliyle kısa "
  "bir amaç açıklaması üretir ve bu birimleri bir vektör veri tabanına (Weaviate) "
  "indeksler. Sorgu anında, kullanıcının sorusu melez arama (BM25 + vektör) ile en "
  "alakalı tablolara eşlenir, tablolar önem katsayısına göre yeniden sıralanır ve "
  "yalnızca seçilen tabloların şeması dil modeline verilir. Üretilen sorgular yalnızca "
  "SELECT ile sınırlandırılarak güvenlik sağlanır; ayrıca soru, yanıtlanabilirliğine "
  "göre üç durumlu bir çözümleyici ile sınıflandırılır. Niteliksel değerlendirme ve "
  "örnek senaryolar, yaklaşımın bağlam bütçesini belirgin biçimde küçülttüğünü ve "
  "ölçeklenebilir, güvenli bir doğal dil veri erişimi sağladığını göstermektedir.")
P("Anahtar Kelimeler: Doğal dilden SQL, büyük dil modelleri, getirim-destekli üretim, "
  "şema bağlama, vektör veri tabanı, mikroservis mimarisi.")

# ── ABSTRACT ──────────────────────────────────────────────────────────────────
T.h5("ABSTRACT")
P("This thesis presents QueryLens, a SaaS platform that lets users connect their own "
  "relational databases through an agent and query them using natural language. The "
  "platform consists of a React front end, a Spring Cloud Gateway, dedicated "
  "authentication, user, database and AI microservices, and a WebSocket agent that "
  "runs inside the user’s own database environment. The core problem addressed is the "
  "central scalability challenge of large-language-model-based text-to-SQL: real "
  "enterprise databases may contain hundreds of tables, so feeding the entire schema "
  "to the model both exceeds context limits and degrades accuracy. This problem is "
  "solved with an adaptation of the Retrieval-Augmented Schema Linking (RASL) "
  "approach. When a database is verified, the system performs automatic introspection "
  "to extract each table’s columns, types, primary/foreign keys and approximate row "
  "count; it generates a short purpose description per table with a large language "
  "model and indexes these units into a vector database (Weaviate). At query time the "
  "user’s question is matched to the most relevant tables via hybrid search (BM25 + "
  "vector), the tables are re-ranked by an importance score, and only the schema of "
  "the selected tables is given to the language model. Generated queries are "
  "restricted to SELECT statements for security, and questions are classified by a "
  "three-case analyzer according to their answerability. Qualitative evaluation and "
  "example scenarios show that the approach substantially reduces the context budget "
  "and provides scalable, secure natural-language data access.")
P("Keywords: Text-to-SQL, large language models, retrieval-augmented generation, "
  "schema linking, vector database, microservice architecture.")


# ══════════════════════════════════════════════════════════════════════════════
# TEZ METNİ BÖLÜMÜ (arap rakamı)
# ══════════════════════════════════════════════════════════════════════════════
T.doc.add_section(WD_SECTION.NEW_PAGE)

# Bu noktadan sonraki içerik gen_thesis_body.py içinde yazılır (okunabilirlik için).
import gen_thesis_body  # noqa: E402
gen_thesis_body.write_body(T)


# ══════════════════════════════════════════════════════════════════════════════
# BÖLÜM / SAYFA NUMARALANDIRMA AYARLARI
# ══════════════════════════════════════════════════════════════════════════════
secs = T.doc.sections
# secs[0] = kapaklar (numarasız), secs[1] = ön sayfalar (roma), secs[2] = tez metni (arap)
cover = secs[0]
front = secs[1]
body = secs[2]

# Kapak: şablondan miras kalan (sayfa numaralı) header/footer referanslarını temizle → numarasız
clear_hf_refs(cover)
clear_pgnum(cover)

# Ön sayfalar: alt-orta roma rakamı, ii'den başlar (iç kapak i sayılır, yazılmaz)
clear_hf_refs(front)
front.footer.is_linked_to_previous = False
set_page_number_format(front, "lowerRoman", start=2)
put_page_field(front.footer)

# Tez metni: ALT-orta arap rakamı, 1'den başlar (kullanıcı isteği: numaralar altta)
clear_hf_refs(body)
body.header.is_linked_to_previous = False   # üst bilgi boş (üstte numara olmasın)
body.footer.is_linked_to_previous = False
set_page_number_format(body, "decimal", start=1)
put_page_field(body.footer)

T.save(OUT)
print("Tez üretildi:", OUT)
